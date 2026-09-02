"""
Document Processor for EduBot — Multi-Format Support.

Supports: PDF, DOCX, TXT

=== ADMIN-ONLY ===
Document upload and management is admin-only.
Students never upload documents — they just ask questions.
Administrators build the knowledge base once, and all students benefit.

=== SUPPORTED FORMATS ===
1. PDF (.pdf) — Extracted with PyPDF
2. DOCX (.docx) — Extracted with python-docx
3. TXT (.txt) — Read directly as text
"""

import shutil
from pathlib import Path
from typing import List, Tuple

from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.database import Document
from app.models.enums import DocumentStatus
from app.services.vector_store import vector_store_service
from app.utils.exceptions import DocumentProcessingError
from app.utils.logger import get_logger

logger = get_logger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentProcessor:
    """
    Processes uploaded documents (PDF, DOCX, TXT) for the RAG pipeline.

    Workflow:
    1. Validate uploaded file (type, size)
    2. Save to disk
    3. Extract text (format-specific)
    4. Split into chunks
    5. Add chunks to FAISS vector store
    6. Update database record
    """

    def __init__(self) -> None:
        self._settings = get_settings()

    async def process_upload(
        self,
        file_content: bytes,
        filename: str,
        session: AsyncSession,
    ) -> Document:
        """
        Process an uploaded document file end-to-end.

        Parameters
        ----------
        file_content : bytes
            The raw file content
        filename : str
            Original filename (e.g., "intro_to_ml.pdf")
        session : AsyncSession
            Database session

        Returns
        -------
        Document
            The database record with processing results
        """
        self._validate_file(file_content, filename)

        document = Document(
            filename=filename,
            status=DocumentStatus.PENDING,
            file_size=len(file_content),
        )
        session.add(document)
        await session.flush()

        try:
            document.status = DocumentStatus.PROCESSING

            file_path = await self._save_file(file_content, filename)
            logger.info(f"Saved uploaded file: {file_path}")

            # Extract text based on file type
            extension = Path(filename).suffix.lower()
            text = self._extract_text(file_path, extension)

            if not text.strip():
                raise DocumentProcessingError(
                    "Document contains no extractable text",
                    filename=filename,
                )

            chunks, metadatas = self._split_into_chunks(text, filename)
            logger.info(f"Split '{filename}' into {len(chunks)} chunks.")

            added_count = await vector_store_service.add_documents(
                texts=chunks,
                metadatas=metadatas,
            )

            document.status = DocumentStatus.COMPLETED
            document.chunk_count = added_count

            logger.info(
                f"Successfully processed '{filename}': "
                f"{added_count} chunks indexed."
            )

            return document

        except DocumentProcessingError:
            document.status = DocumentStatus.FAILED
            document.error_message = "Processing failed"
            raise

        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.error_message = str(e)
            logger.error(f"Failed to process '{filename}': {e}")
            raise DocumentProcessingError(
                f"Processing failed: {e}",
                filename=filename,
            ) from e

    def _validate_file(self, file_content: bytes, filename: str) -> None:
        """Validate the uploaded file."""
        if not file_content:
            raise DocumentProcessingError("File is empty", filename=filename)

        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise DocumentProcessingError(
                f"Unsupported file type: {extension}. "
                f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}",
                filename=filename,
            )

        if len(file_content) > self._settings.MAX_UPLOAD_SIZE:
            max_mb = self._settings.MAX_UPLOAD_SIZE / (1024 * 1024)
            raise DocumentProcessingError(
                f"File exceeds maximum size of {max_mb:.0f} MB",
                filename=filename,
            )

    async def _save_file(self, file_content: bytes, filename: str) -> Path:
        """Save the uploaded file to the uploads directory."""
        upload_dir = self._settings.upload_dir_path
        upload_dir.mkdir(parents=True, exist_ok=True)

        file_path = upload_dir / filename

        counter = 1
        while file_path.exists():
            stem = Path(filename).stem
            suffix = Path(filename).suffix
            file_path = upload_dir / f"{stem}_{counter}{suffix}"
            counter += 1

        file_path.write_bytes(file_content)
        return file_path

    def _extract_text(self, file_path: Path, extension: str) -> str:
        """
        Extract text from a document based on its file type.

        Dispatches to the appropriate extraction method.
        """
        if extension == ".pdf":
            return self._extract_text_pdf(file_path)
        elif extension == ".docx":
            return self._extract_text_docx(file_path)
        elif extension == ".txt":
            return self._extract_text_txt(file_path)
        else:
            raise DocumentProcessingError(
                f"No text extractor for: {extension}",
                filename=file_path.name,
            )

    def _extract_text_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(str(file_path))
            text_parts: List[str] = []

            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n[PAGE {page_num}]\n{page_text}")

            full_text = "\n".join(text_parts)
            logger.debug(
                f"Extracted {len(full_text)} chars from {len(reader.pages)} PDF pages."
            )
            return full_text

        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract text from PDF: {e}",
                filename=file_path.name,
            ) from e

    def _extract_text_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file.

        Uses python-docx to read paragraphs from Word documents.
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            text_parts: List[str] = []

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            full_text = "\n".join(text_parts)
            logger.debug(
                f"Extracted {len(full_text)} chars from DOCX "
                f"({len(doc.paragraphs)} paragraphs)."
            )
            return full_text

        except ImportError:
            raise DocumentProcessingError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install python-docx",
                filename=file_path.name,
            )
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract text from DOCX: {e}",
                filename=file_path.name,
            ) from e

    def _extract_text_txt(self, file_path: Path) -> str:
        """Extract text from a plain text file."""
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            logger.debug(f"Read {len(text)} chars from TXT file.")
            return text

        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to read text file: {e}",
                filename=file_path.name,
            ) from e

    def _split_into_chunks(
        self,
        text: str,
        filename: str,
    ) -> Tuple[List[str], List[dict]]:
        """Split extracted text into overlapping chunks."""
        chunk_size = self._settings.CHUNK_SIZE
        chunk_overlap = self._settings.CHUNK_OVERLAP

        chunks: List[str] = []
        metadatas: List[dict] = []

        step = chunk_size - chunk_overlap
        if step <= 0:
            step = chunk_size

        current_page = 1

        for i in range(0, len(text), step):
            chunk = text[i : i + chunk_size].strip()

            if not chunk:
                continue

            page_marker_pos = text.rfind("[PAGE ", 0, i + chunk_size)
            if page_marker_pos != -1:
                try:
                    page_end = text.index("]", page_marker_pos)
                    page_str = text[page_marker_pos + 6 : page_end]
                    current_page = int(page_str)
                except (ValueError, IndexError):
                    pass

            chunks.append(chunk)
            metadatas.append({
                "source": filename,
                "page": current_page,
                "chunk_index": len(chunks) - 1,
            })

        return chunks, metadatas

    async def get_all_documents(self, session: AsyncSession) -> List[Document]:
        """Get all documents from the database."""
        result = await session.execute(
            select(Document).order_by(Document.created_at.desc())
        )
        return list(result.scalars().all())

    async def delete_document(
        self,
        document_id: int,
        session: AsyncSession,
    ) -> bool:
        """
        Delete a document from the database.

        Note: This removes the DB record but does NOT remove chunks
        from FAISS (FAISS doesn't support single-document deletion easily).
        Use rebuild_index() after deleting documents to rebuild FAISS.
        """
        result = await session.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalar_one_or_none()

        if not document:
            return False

        # Delete the file from disk if it exists
        file_path = self._settings.upload_dir_path / document.filename
        if file_path.exists():
            file_path.unlink()
            logger.info(f"Deleted file: {file_path}")

        await session.delete(document)
        logger.info(f"Deleted document record: {document.filename}")
        return True

    async def rebuild_index(self, session: AsyncSession) -> int:
        """
        Rebuild the FAISS index from all completed documents.

        This is needed after deleting documents, since FAISS
        doesn't support selective deletion.

        Returns the total number of chunks indexed.
        """
        # Clear existing FAISS index
        await vector_store_service.delete_all()

        # Get all completed documents
        result = await session.execute(
            select(Document).where(Document.status == DocumentStatus.COMPLETED)
        )
        documents = list(result.scalars().all())

        if not documents:
            logger.info("No documents to rebuild index from.")
            return 0

        total_chunks = 0
        for doc in documents:
            file_path = self._settings.upload_dir_path / doc.filename
            if not file_path.exists():
                logger.warning(f"File not found for reindex: {doc.filename}")
                continue

            try:
                extension = Path(doc.filename).suffix.lower()
                text = self._extract_text(file_path, extension)
                if not text.strip():
                    continue

                chunks, metadatas = self._split_into_chunks(text, doc.filename)
                added = await vector_store_service.add_documents(
                    texts=chunks,
                    metadatas=metadatas,
                )
                total_chunks += added
                logger.info(f"Re-indexed '{doc.filename}': {added} chunks")

            except Exception as e:
                logger.error(f"Failed to re-index '{doc.filename}': {e}")

        logger.info(f"Index rebuild complete. Total chunks: {total_chunks}")
        return total_chunks


# ============================================
# SINGLETON INSTANCE
# ============================================
document_processor = DocumentProcessor()
